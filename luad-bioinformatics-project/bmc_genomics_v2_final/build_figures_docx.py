#!/usr/bin/env python3
"""Assemble all main + supplementary figures into one Word document with legends."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE=Path("/home/user/-----/luad-bioinformatics-project/bmc_genomics_v2_final")
FIG=BASE/"figures"
OUT=BASE/"LUAD_PLK1_ERO1A_figures.docx"
INK=RGBColor(0x0b,0x0b,0x0b); MUT=RGBColor(0x52,0x51,0x4e)

TITLE=("A literature-guided public-data evidence hierarchy prioritises a "
       "PLK1-centred mitotic/redox programme in lung adenocarcinoma")

LEG=[
("Figure1","Figure 1. Literature-guided TCGA discovery of a PLK1-centred mitotic/redox programme.",
 "(a) Genome-wide TCGA-LUAD survival screen; the top 15 genes are ranked by -log10 log-rank P, with the lead "
 "candidates PLK1 (blue) and ERO1A (orange) highlighted. (b) Cross-layer evidence-support matrix for the nine "
 "programme genes across six independent evidence layers (TCGA FDR significance, TCGA-train log-rank, GSE31210 and "
 "GSE50081 replication, DepMap dependency, spatial support); the right-hand column gives the per-gene support count."),
("Figure2","Figure 2. A nine-gene score: construction, held-out separation and multi-cohort transfer.",
 "(a) Frozen Cox coefficients of the nine programme genes fitted in the TCGA training partition (red, risk-increasing; "
 "blue, protective). (b) Harrell C-index of the prior ridge scaffold versus the consensus score on TCGA training and "
 "held-out data; the dashed line marks random discrimination (0.5). (c) Forest plot of hazard ratio per 1-SD score "
 "(95% CI) on external transfer; the score is significant in four of five cohorts (blue) and does not replicate in "
 "GSE37745 (orange), which is retained as a transparent null. C-index and Cox P values are annotated."),
("Figure3","Figure 3. Pathway and Human Protein Atlas biological-context coherence.",
 "(a) g:Profiler enrichment (-log10 adjusted P, g:SCS correction) of the candidate programme, coloured by theme "
 "(blue, mitotic/checkpoint; orange, ER-redox/folding). (b) Distribution of Human Protein Atlas pathology-direction "
 "concordance across evaluated candidate genes."),
("Figure4","Figure 4. Single-cell aggregate and Visium tissue-context evidence.",
 "(a) E-MTAB-13530 Visium tumour-minus-adjacent candidate-gene expression (delta log1p CP10K); ERO1A (orange) and "
 "PLK1 (blue) are highlighted. (b) Spot-level Spearman correlation of each programme axis with the tumour-like state, "
 "the cross-method-robust spatial signal; positive in all 12 tumour sections."),
("Figure5","Figure 5. Functional prioritisation of PLK1 with explicit therapeutic boundaries.",
 "(a) DepMap Public 26Q1 dependency ranking (mean LUAD CRISPR effect); PLK1 (blue) is the strongest dependency, past "
 "the -0.5 common-dependency threshold. (b) PLK1 mean CRISPR effect across LUAD, other-lung and non-lung models, "
 "showing common essentiality rather than LUAD selectivity (Mann-Whitney P=0.73). (c) Expression-drug associations "
 "(Spearman r versus -log10 FDR) for PLK1 (blue) and ERO1A (orange); none of the 12 tests passes FDR<0.05 (dashed "
 "line), so no PLK-inhibitor stratifier is supported."),
("Figure6_singlecell","Figure 6. Single-cell (GSE131907, 208,506 cells): the programme is enriched in malignant/tumour-derived epithelium.",
 "(a) Programme-gene expression across annotated cell types (dot size = percent expressing, colour = mean expression). (b) Malignant (tumour-origin) versus normal epithelial differential expression; all nine genes are enriched at the sample level (36 tumour-origin vs 11 normal-lung samples, FDR<0.05), avoiding cell-level pseudoreplication. (c) Programme score peaks in malignant and tumour-transitional epithelial states versus normal epithelium. (d) Programme and PLK1 correlate with proliferation (MKI67) across epithelial cells."),
("Figure7_spatial_deconv","Figure 7. Reference-based Visium deconvolution (programme genes excluded from the signature; 22 sections, five patients).",
 "(a) Spearman correlation between programme score and inferred cell-type abundance across ~53,000 spots; the programme co-localises with the epithelial compartment and negatively with mast/endothelial cells. (b) With programme genes excluded from the signature, the programme-versus-epithelial correlation is positive in 21/22 sections and 5/5 patients; the patient-level test is under-powered (n=5)."),
("Figure8_nomogram","Figure 8. Incremental prognostic assessment (TCGA-LUAD; train-fit, test-evaluated).",
 "(a) Harrell C-index for the gene score alone versus the score-plus-clinical nomogram (all TCGA and held-out test); the score is independent (LR p<1e-8) but adds no discrimination over clinical staging (clinical 0.72, gene 0.64, combined 0.73). (b) Incremental delta-C-index over the clinical model is +0.006 (95% CI crosses 0). (c) Three-year calibration over-predicts risk."),
("FigureS1","Additional file 1 (Supplementary Figure S1). iLINCS/L1000 PLK1-ERO1A perturbation convergence.",
 "(a) Signature-level Spearman correlation (978 shared L1000 genes): within-cell PLK1-knockdown vs ERO1A-knockdown "
 "convergence in HCC515 and A549, and cross-cell knockdown reproducibility. (b) Interpretation boundary: convergence "
 "appears only in HCC515, the PLK1 knockdown signature is not reproducible across cell lines, and signature overlap "
 "does not establish molecular interaction, epistasis, pathway ordering or drug synergy."),
("FigureS2","Additional file 2 (Supplementary Figure S2). PLK1 lacks a demonstrated LUAD-selective therapeutic window.",
 "(a) PLK1 CRISPR dependency fraction in LUAD versus non-cancer models: pan-essential (1.00) in both. (b) Coverage "
 "audit: no non-cancerous lung model has PLK1 CRISPR or PLK-inhibitor coverage, and LUAD versus non-cancer PLK1/ERO1A "
 "expression differences are non-significant; no therapeutic-window or normal-tissue safety claim is supported."),
]

doc=Document()
for sec in doc.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)
st=doc.styles["Normal"].font; st.name="Calibri"; st.size=Pt(10.5)

h=doc.add_paragraph(); r=h.add_run("Figures"); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=INK
p=doc.add_paragraph(); r=p.add_run(TITLE); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=MUT
p=doc.add_paragraph(); r=p.add_run("Target journal: BMC Genomics · 5 main figures + 2 additional-file (supplementary) figures. "
    "Figures are 400-dpi renderings; editable vector PDFs are provided in the submission package.")
r.font.size=Pt(9); r.font.color.rgb=MUT

for stem,cap,body in LEG:
    doc.add_paragraph()
    png=FIG/f"{stem}.png"
    ip=doc.add_paragraph(); ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(str(png), width=Cm(16.5))
    cp=doc.add_paragraph()
    cr=cp.add_run(cap); cr.bold=True; cr.font.size=Pt(10); cr.font.color.rgb=INK
    lp=doc.add_paragraph()
    lr=lp.add_run(" "+body); lr.font.size=Pt(9.5); lr.font.color.rgb=MUT

doc.save(str(OUT))
print("wrote", OUT)
