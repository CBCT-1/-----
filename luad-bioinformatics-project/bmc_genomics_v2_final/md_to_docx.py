#!/usr/bin/env python3
"""Minimal, clean Markdown -> DOCX for the BMC manuscript (headings, tables, paragraphs, refs)."""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

src=Path(sys.argv[1]); dst=Path(sys.argv[2])
lines=src.read_text(encoding="utf-8").splitlines()
INK=RGBColor(0x0b,0x0b,0x0b)

doc=Document()
for s in doc.sections:
    s.top_margin=Cm(2.2); s.bottom_margin=Cm(2.2); s.left_margin=Cm(2.4); s.right_margin=Cm(2.4)
f=doc.styles["Normal"].font; f.name="Calibri"; f.size=Pt(11)

def emit_table(block):
    rows=[[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
    rows=[r for r in rows if not all(set(c)<=set("-: ") for c in r)]  # drop separator row
    if not rows: return
    t=doc.add_table(rows=len(rows), cols=len(rows[0])); t.style="Light Grid Accent 1"
    t.autofit=True
    for i,r in enumerate(rows):
        for j,c in enumerate(r):
            if j<len(t.rows[i].cells):
                cell=t.rows[i].cells[j]; cell.text=""
                run=cell.paragraphs[0].add_run(re.sub(r"\*\*(.+?)\*\*",r"\1",c))
                run.font.size=Pt(8.5)
                if i==0: run.bold=True

i=0; tbl=[]
def flush_tbl():
    global tbl
    if tbl: emit_table(tbl); tbl=[]
while i<len(lines):
    ln=lines[i].rstrip()
    if ln.startswith("|"):
        tbl.append(ln); i+=1; continue
    flush_tbl()
    if not ln.strip(): i+=1; continue
    if ln.startswith("# "):
        p=doc.add_paragraph(); r=p.add_run(ln[2:]); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=INK
    elif ln.startswith("### "):
        p=doc.add_paragraph(); r=p.add_run(ln[4:]); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=INK
    elif ln.startswith("## "):
        p=doc.add_paragraph(); r=p.add_run(ln[3:]); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=INK
    elif ln.startswith("---"):
        pass
    elif re.match(r"^\d+\.\s", ln):  # numbered reference
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.8); p.paragraph_format.first_line_indent=Cm(-0.8)
        r=p.add_run(ln); r.font.size=Pt(9.5)
    else:
        txt=re.sub(r"\*\*(.+?)\*\*",r"\1",ln); txt=re.sub(r"\*(.+?)\*",r"\1",txt)
        p=doc.add_paragraph(); p.add_run(txt)
    i+=1
flush_tbl()
doc.save(str(dst)); print("wrote",dst)
